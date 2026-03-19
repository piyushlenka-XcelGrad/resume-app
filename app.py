import streamlit as st
import os
from io import BytesIO
import pandas as pd
import re
import PyPDF2
from typing import List, Dict, Tuple, Optional
import datetime
from docx import Document
import requests


# ═══════════════════════════════════════════════════════════════════
# CONFIGURATION — paste your Google Drive API key here
# ═══════════════════════════════════════════════════════════════════

GOOGLE_DRIVE_API_KEY = "AIzaSyB33hUEZClIvP662hWsdqGCDcaqTz4zA5I"


# ═══════════════════════════════════════════════════════════════════
# GOOGLE DRIVE HELPERS
# ═══════════════════════════════════════════════════════════════════

def parse_drive_id(url_or_id: str) -> Tuple[Optional[str], str]:
    s = url_or_id.strip()
    m = re.search(r'/folders/([a-zA-Z0-9_-]{25,})', s)
    if m:
        return m.group(1), 'folder'
    m = re.search(r'/file/d/([a-zA-Z0-9_-]{25,})', s)
    if m:
        return m.group(1), 'file'
    m = re.search(r'[?&]id=([a-zA-Z0-9_-]{25,})', s)
    if m:
        return m.group(1), 'folder'
    m = re.fullmatch(r'[a-zA-Z0-9_-]{25,}', s)
    if m:
        return s, 'folder'
    return None, 'unknown'


def list_drive_folder(folder_id: str, api_key: str) -> List[Dict]:
    """List all PDF/DOCX files in a public Google Drive folder."""
    files = []
    page_token = None
    allowed_mimes = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    }
    while True:
        params = {
            'key': api_key,
            'q': f"'{folder_id}' in parents and trashed=false",
            'fields': 'nextPageToken, files(id, name, mimeType, webContentLink)',
            'pageSize': 100,
        }
        if page_token:
            params['pageToken'] = page_token
        resp = requests.get(
            'https://www.googleapis.com/drive/v3/files',
            params=params,
            timeout=20,
        )
        resp.raise_for_status()
        data = resp.json()
        for f in data.get('files', []):
            if f.get('mimeType') in allowed_mimes:
                files.append(f)
        page_token = data.get('nextPageToken')
        if not page_token:
            break
    return files


def download_drive_file_public(file_id: str) -> bytes:
    """
    Download a publicly shared Google Drive file using the
    public export URL — no API key or OAuth needed.
    Handles Google's virus-scan confirmation page automatically.
    """
    session = requests.Session()

    # Primary URL — works for most files
    url = f"https://drive.google.com/uc?export=download&id={file_id}"
    resp = session.get(url, timeout=60)
    resp.raise_for_status()

    # If Google returns a confirmation page (large files / virus scan)
    # extract the confirm token and retry
    if b'confirm=' in resp.content or b'virus scan warning' in resp.content.lower():
        # Try to find confirm token
        match = re.search(rb'confirm=([0-9A-Za-z_\-]+)', resp.content)
        if match:
            confirm = match.group(1).decode()
            url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm={confirm}"
            resp = session.get(url, timeout=60)
            resp.raise_for_status()

    # Check we actually got a file (not an HTML error page)
    content_type = resp.headers.get('Content-Type', '')
    if 'text/html' in content_type and len(resp.content) < 50_000:
        raise ValueError(
            "Got an HTML page instead of a file. "
            "Make sure the file is shared as 'Anyone with the link'."
        )

    return resp.content


def fetch_resumes_from_drive(
    folder_id: str,
    api_key: str,
    status_placeholder,
    progress_bar,
) -> Tuple[List[Dict], List[str]]:
    raw: List[Dict]   = []
    errors: List[str] = []

    # Step 1: list files via API
    status_placeholder.text("📂 Fetching file list from Google Drive…")
    try:
        drive_files = list_drive_folder(folder_id, api_key)
    except requests.HTTPError as e:
        code = e.response.status_code if e.response is not None else '?'
        if code == 403:
            errors.append(
                "Access denied (403). Make sure the folder is shared as "
                "'Anyone with the link can view'."
            )
        elif code == 404:
            errors.append("Folder not found (404). Check the folder link.")
        else:
            errors.append(f"Drive API error {code}: {e}")
        return raw, errors
    except Exception as e:
        errors.append(f"Could not reach Google Drive: {e}")
        return raw, errors

    if not drive_files:
        errors.append("No PDF or DOCX files found in the folder.")
        return raw, errors

    status_placeholder.text(f"📂 Found {len(drive_files)} file(s). Downloading…")

    # Step 2: download each file via public URL (no API key needed)
    for i, f in enumerate(drive_files):
        try:
            status_placeholder.text(
                f"⬇️  Downloading {i + 1}/{len(drive_files)}: {f['name']}"
            )
            file_bytes = download_drive_file_public(f['id'])
            raw.append({'filename': f['name'], 'bytes': file_bytes})
        except Exception as e:
            errors.append(f"{f['name']}: {e}")
        progress_bar.progress((i + 1) / len(drive_files))

    return raw, errors


# ═══════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    try:
        doc = Document(BytesIO(docx_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text(filename: str, file_bytes: bytes) -> str:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_text_from_pdf_bytes(file_bytes)
    elif fname.endswith(".docx"):
        return extract_text_from_docx_bytes(file_bytes)
    return ""


# ═══════════════════════════════════════════════════════════════════
# FIELD EXTRACTION
# ═══════════════════════════════════════════════════════════════════

def extract_name_from_filename(filename: str) -> str:
    name = os.path.splitext(filename)[0]
    name = name.replace('_', ' ').replace('-', ' ')
    return ' '.join(w.capitalize() for w in name.split() if w) or filename


def extract_email(text: str) -> str:
    if not text:
        return ""
    clean = ' '.join(text.split())
    patterns = [
        r'\b([a-zA-Z0-9][a-zA-Z0-9._%+\-]*@[a-zA-Z0-9][a-zA-Z0-9.\-]*\.[a-zA-Z]{2,})\b',
        r'([a-zA-Z0-9._%+\-]+)\s*@\s*([a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
        r'(?:email|e-mail|mail)\s*[:\-]?\s*([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
    ]
    candidates = []
    for pat in patterns:
        for m in re.finditer(pat, clean, re.IGNORECASE):
            if len(m.groups()) == 1:
                candidates.append(m.group(1).lower().strip('.,;:'))
            elif len(m.groups()) == 2:
                candidates.append(f"{m.group(1)}@{m.group(2)}".lower())
    false_positives = {'example.com', 'test.com', 'domain.com', 'email.com', 'sample.com'}
    for email in candidates:
        parts = email.split('@')
        if len(parts) == 2:
            user, domain = parts
            if (re.match(r'^[a-z0-9._%+\-]+$', user)
                    and '.' in domain
                    and len(domain.split('.')[-1]) >= 2
                    and email not in false_positives
                    and len(user) >= 2):
                return email
    return ""


def extract_phone(text: str) -> str:
    if not text:
        return ""
    norm = ' '.join(text.split())
    patterns = [
        r'\+91[\s\-]?[6-9]\d{9}',
        r'\b[6-9]\d{4}[\s\-]\d{5}\b',
        r'\b[6-9]\d{9}\b',
        r'\+\d{1,3}[\s\-]?\d{10,12}',
        r'\b\d{3}[\s.\-]\d{3}[\s.\-]\d{4}\b',
        r'\b\d{5}[\s\-]\d{5}\b',
        r'(?:phone|mobile|cell|contact|tel)\s*[:\-]?\s*([\+\d][\d\s\-\(\)]{9,})',
    ]
    for pat in patterns:
        m = re.search(pat, norm, re.IGNORECASE)
        if m:
            phone = (m.group(1) if m.lastindex and m.lastindex >= 1 and re.search(r'phone|mobile|cell|contact|tel', pat, re.I) else m.group(0)).strip()
            if len(re.sub(r'\D', '', phone)) >= 10:
                return phone
    return ""


def extract_education(text: str) -> str:
    if not text:
        return ""
    edu_start = -1
    for kw in [r'\beducation\b', r'\bacademics?\b', r'\bqualifications?\b', r'\beducational\s+background\b']:
        m = re.search(kw, text, re.IGNORECASE)
        if m:
            edu_start = m.start()
            break
    search_text = text[edu_start:edu_start + 1000] if edu_start != -1 else text[:1200]
    degree_patterns = [
        r'(Bachelor\s+of\s+[A-Za-z\s]{2,40})',
        r'(Master\s+of\s+[A-Za-z\s]{2,40})',
        r'(B\.?\s*Tech\.?(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(M\.?\s*Tech\.?(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(B\.?E\.?(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(M\.?E\.?(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(B\.?Sc\.?(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(M\.?Sc\.?(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(BCA(?:\s+in\s+[A-Za-z\s]{2,25})?)',
        r'(MCA(?:\s+in\s+[A-Za-z\s]{2,25})?)',
        r'(MBA(?:\s+in\s+[A-Za-z\s&]{2,35})?)',
        r'(Ph\.?D\.?(?:\s+in\s+[A-Za-z\s]{2,35})?)',
        r'(Diploma\s+in\s+[A-Za-z\s&]{2,40})',
        r'(Intermediate|Higher\s+Secondary|10\+2|HSC|SSC)',
    ]
    matches = []
    for pat in degree_patterns:
        for m in re.finditer(pat, search_text, re.IGNORECASE):
            matches.append(m.group(1).strip())
    if matches:
        best = max(matches, key=len)
        return re.sub(r'\s+', ' ', best.strip())[:160]
    return ""


def extract_location(text: str) -> str:
    if not text:
        return ""
    top = text[:1000]
    for pat in [
        r'(?:location|address|city|based\s+in|residence|residing\s+at|current\s+address)\s*[:\-]?\s*([A-Z][^\n]{3,55})',
    ]:
        m = re.search(pat, top, re.IGNORECASE)
        if m:
            loc = m.group(1).strip().split('\n')[0].split('|')[0]
            if '@' not in loc and not re.search(r'\d{10}', loc) and len(loc) < 65:
                return re.sub(r'\s+', ' ', loc).strip('.,')
    m = re.search(
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b',
        top
    )
    if m and not re.search(
        r'University|College|Institute|School|Ltd|Inc|Pvt|Corp|Solutions',
        m.group(0), re.IGNORECASE
    ):
        return m.group(0)
    cities = [
        'Navi Mumbai', 'Greater Noida', 'New Delhi',
        'Mumbai', 'Delhi', 'Bangalore', 'Bengaluru', 'Hyderabad', 'Chennai',
        'Kolkata', 'Pune', 'Ahmedabad', 'Noida', 'Gurgaon', 'Gurugram',
        'Thane', 'Indore', 'Jaipur', 'Lucknow', 'Bhopal', 'Surat', 'Nagpur',
        'Panvel', 'Vadodara', 'Coimbatore', 'Kochi', 'Visakhapatnam',
        'Patna', 'Ranchi', 'Chandigarh', 'Mysuru', 'Mysore', 'Bhubaneswar'
    ]
    for city in cities:
        if re.search(r'\b' + re.escape(city) + r'\b', top, re.IGNORECASE):
            return city
    return ""


def is_internship(block: str) -> bool:
    return bool(re.search(
        r'\bintern(?:ship)?\b|\btrainee\b|\bapprentice(?:ship)?\b'
        r'|\bindustrial\s+training\b|\bpractical\s+training\b|\bvocational\b',
        block, re.IGNORECASE
    ))


def extract_total_experience(text: str) -> float:
    if not text:
        return 0.0
    now = datetime.datetime.now()
    cy, cm = now.year, now.month
    month_map = {
        'jan': 1, 'january': 1, 'feb': 2, 'february': 2,
        'mar': 3, 'march': 3, 'apr': 4, 'april': 4,
        'may': 5, 'jun': 6, 'june': 6, 'jul': 7, 'july': 7,
        'aug': 8, 'august': 8, 'sep': 9, 'sept': 9, 'september': 9,
        'oct': 10, 'october': 10, 'nov': 11, 'november': 11,
        'dec': 12, 'december': 12
    }
    text_lower = text.lower()
    for pat in [
        r'(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)[\s,]+(?:of\s+)?(?:total\s+)?(?:work\s+|professional\s+)?experience',
        r'(?:total\s+)?(?:work\s+)?experience\s*[:\-]?\s*(\d+\.?\d*)\s*(?:years?|yrs?)',
        r'(\d+)\s*\+?\s*years?\s+(?:of\s+)?(?:work\s+|professional\s+)?experience',
    ]:
        for m in re.finditer(pat, text_lower):
            ctx = text_lower[max(0, m.start() - 150): m.end() + 50]
            if not is_internship(ctx):
                try:
                    val = float(m.group(1))
                    if 0 < val <= 50:
                        return round(val, 1)
                except Exception:
                    pass
    entries = re.split(r'\n\s*\n+', text)
    total_months = 0
    date_pats = [
        (r'([a-z]{3,9})\s+(\d{4})\s*[-–—to]+\s*(present|current|now|ongoing|till\s*date)', 'mypres'),
        (r'([a-z]{3,9})\s+(\d{4})\s*[-–—to]+\s*([a-z]{3,9})\s+(\d{4})', 'mymy'),
        (r'(\d{1,2})/(\d{4})\s*[-–—to]+\s*(present|current|now)', 'mmpres'),
        (r'(\d{1,2})/(\d{4})\s*[-–—to]+\s*(\d{1,2})/(\d{4})', 'mmyymmy'),
        (r'(\d{4})\s*[-–—to]+\s*(present|current|now|ongoing|till\s*date)', 'ypres'),
        (r'(\d{4})\s*[-–—to]+\s*(\d{4})', 'yy'),
    ]
    for entry in entries:
        if is_internship(entry):
            continue
        el = entry.lower()
        for pat, pt in date_pats:
            matched_range = False
            for m in re.finditer(pat, el):
                try:
                    sm = em = 1
                    sy = ey = 0
                    if pt == 'mypres':
                        sm = month_map.get(m.group(1)[:3], 1); sy = int(m.group(2))
                        em = cm; ey = cy
                    elif pt == 'mymy':
                        sm = month_map.get(m.group(1)[:3], 1); sy = int(m.group(2))
                        em = month_map.get(m.group(3)[:3], 12); ey = int(m.group(4))
                    elif pt == 'mmpres':
                        sm = int(m.group(1)); sy = int(m.group(2)); em = cm; ey = cy
                    elif pt == 'mmyymmy':
                        sm = int(m.group(1)); sy = int(m.group(2))
                        em = int(m.group(3)); ey = int(m.group(4))
                    elif pt == 'ypres':
                        sm = 1; sy = int(m.group(1)); em = cm; ey = cy
                    elif pt == 'yy':
                        sm = 1; sy = int(m.group(1)); em = 12; ey = int(m.group(2))
                    if 1970 <= sy <= cy and sy <= ey <= cy + 1:
                        months = (ey - sy) * 12 + (em - sm)
                        if 0 < months <= 600:
                            total_months += months
                            matched_range = True
                            break
                except Exception:
                    pass
            if matched_range:
                break
    if total_months > 0:
        return round(min(total_months / 12, 50), 1)
    return 0.0


def extract_companies(text: str) -> List[str]:
    if not text:
        return []
    exp_start = -1
    for kw in [r'\bwork\s+experience\b', r'\bprofessional\s+experience\b',
                r'\bexperience\b', r'\bemployment\b', r'\bwork\s+history\b']:
        m = re.search(kw, text, re.IGNORECASE)
        if m:
            exp_start = m.start()
            break
    search_text = text[exp_start:] if exp_start != -1 else text
    companies = []
    suffix = (
        r'(?:Pvt\.?\s*Ltd\.?|Private\s+Limited|Ltd\.?|Limited|Inc\.?|'
        r'Corp(?:oration)?|LLP|Technologies|Technology|Tech|Solutions|'
        r'Services|Systems|Software|Digital|Analytics|Consulting|Consultancy|'
        r'Group|Holdings|Enterprises|International|Global|Networks|'
        r'Infotech|InfoSystems|Associates|Ventures|Industries|'
        r'Communications|Labs?|Media|Innovations?|Platforms?|Products?|'
        r'Outsourcing|Recruitment|Staffing)'
    )
    for m in re.finditer(rf'([A-Z][A-Za-z0-9\s&\.,\-]{{1,45}}\s+{suffix})', search_text):
        c = re.sub(r'\s+', ' ', m.group(1).strip())
        if 4 < len(c) < 75:
            companies.append(c)
    for m in re.finditer(
        r'(?:Company|Employer|Organization|Firm|Client|Worked\s+at|Working\s+at)\s*[:\-]\s*([^\n,]{3,65})',
        search_text, re.IGNORECASE
    ):
        companies.append(m.group(1).strip())
    for line in search_text.split('\n'):
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            first = parts[0]
            if (first
                    and re.match(r'^[A-Z]', first)
                    and 4 < len(first) < 65
                    and not re.match(
                        r'^(MANAGER|DEVELOPER|ENGINEER|ANALYST|INTERN|TRAINEE|ASSOCIATE|SENIOR|JUNIOR)',
                        first.upper()
                    )):
                companies.append(first)
    for m in re.finditer(rf'\bat\s+([A-Z][A-Za-z0-9\s&]+(?:{suffix}))', search_text):
        c = re.sub(r'\s+', ' ', m.group(1).strip())
        if 4 < len(c) < 75:
            companies.append(c)
    seen: set = set()
    unique: List[str] = []
    for c in companies:
        key = re.sub(r'\s+', ' ', c.lower().strip())
        if key not in seen and len(key) > 3:
            seen.add(key)
            unique.append(c.strip())
    return unique[:8]


# ═══════════════════════════════════════════════════════════════════
# PROCESS SINGLE RESUME
# ═══════════════════════════════════════════════════════════════════

def process_resume(file_bytes: bytes, filename: str) -> Optional[Dict]:
    full_text = extract_text(filename, file_bytes)
    if not full_text.strip():
        return None
    companies = extract_companies(full_text)
    return {
        'Filename':          filename,
        'Name':              extract_name_from_filename(filename),
        'Email':             extract_email(full_text),
        'Phone':             extract_phone(full_text),
        'Education':         extract_education(full_text),
        'Location':          extract_location(full_text),
        'Total Exp (Years)': extract_total_experience(full_text),
        'Companies':         ', '.join(companies) if companies else '',
        '_full_text':        full_text.lower(),
    }


# ═══════════════════════════════════════════════════════════════════
# SEARCH / FILTER
# ═══════════════════════════════════════════════════════════════════

def search_resumes(
    all_data: List[Dict],
    terms: List[str],
    match_all: bool = False
) -> Tuple[List[Dict], List[Dict]]:
    matching:     List[Dict] = []
    non_matching: List[Dict] = []
    clean_terms = [t.strip() for t in terms if t.strip()]
    for resume in all_data:
        ft = resume.get('_full_text', '')
        matched = []
        for term in clean_terms:
            pattern = r'(?<![a-zA-Z0-9])' + re.escape(term.lower()) + r'(?![a-zA-Z0-9])'
            if re.search(pattern, ft):
                matched.append(term)
        is_match = (len(matched) == len(clean_terms)) if match_all else (len(matched) > 0)
        display = {k: v for k, v in resume.items() if k != '_full_text'}
        display['Matched Terms'] = ', '.join(matched) if matched else '—'
        if is_match:
            matching.append(display)
        else:
            non_matching.append(display)
    return matching, non_matching


# ═══════════════════════════════════════════════════════════════════
# EXCEL HELPER
# ═══════════════════════════════════════════════════════════════════

def to_excel(data: List[Dict]) -> bytes:
    df = pd.DataFrame(data)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Resumes')
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="Resume Extractor & Smart Search",
    layout="wide",
    page_icon="📄"
)

# ── Session state initialisation ─────────────────────────────────
_defaults = {
    'processed_data':  [],
    'processing_done': False,
    'matching':        [],
    'non_matching':    [],
    'searched':        False,
    'last_terms':      '',
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Page header ───────────────────────────────────────────────────
st.title("📄 Resume Extractor & Smart Search")
st.markdown(
    "Upload resumes **or paste a Google Drive folder link** → "
    "auto-extract key info → **dynamically search** by skill, company, location, or any keyword"
)

if st.session_state.processing_done:
    if st.button("🔄 Reset — Upload New Files"):
        for k, v in _defaults.items():
            st.session_state[k] = v
        st.rerun()

st.divider()

# ════════════════════════════════════════════════════════════════════
# STEP 1 — SOURCE SELECTION
# ════════════════════════════════════════════════════════════════════
st.header("① Upload & Process Resumes")

if not st.session_state.processing_done:

    source_tab, drive_tab = st.tabs(["📁 Upload Files", "☁️ Google Drive Folder"])

    # ── TAB 1: Local upload ───────────────────────────────────────
    with source_tab:
        uploaded_files = st.file_uploader(
            "Upload Resumes (PDF or DOCX) — up to 100 files",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key="file_uploader",
        )
        if uploaded_files:
            if len(uploaded_files) > 100:
                st.warning("⚠️ Only the first 100 files will be processed.")
                uploaded_files = uploaded_files[:100]
            st.info(f"📂 **{len(uploaded_files)}** file(s) ready")

        if st.button("🚀 Process Uploaded Resumes", type="primary", disabled=not uploaded_files, key="btn_upload"):
            all_data: List[Dict] = []
            failed:   List[str]  = []
            prog   = st.progress(0)
            status = st.empty()
            for i, f in enumerate(uploaded_files):
                status.text(f"Processing {i + 1}/{len(uploaded_files)}: {f.name}")
                try:
                    result = process_resume(f.read(), f.name)
                    if result:
                        all_data.append(result)
                    else:
                        failed.append(f.name)
                except Exception as e:
                    failed.append(f"{f.name} ({e})")
                prog.progress((i + 1) / len(uploaded_files))
            prog.empty(); status.empty()
            if failed:
                st.warning(f"⚠️ Skipped {len(failed)} file(s): " + ", ".join(failed[:5]))
            st.session_state.processed_data  = all_data
            st.session_state.processing_done = True
            st.success(f"✅ Processed **{len(all_data)}** resumes!")
            st.rerun()

    # ── TAB 2: Google Drive ───────────────────────────────────────
    with drive_tab:
        st.markdown(
            """
            **How to use:**
            1. Open the Google Drive folder that contains your resumes
            2. Click **Share → Anyone with the link → Viewer**
            3. Copy the link and paste it below
            """
        )

        drive_link = st.text_input(
            "🔗 Google Drive Folder Link or Folder ID",
            placeholder="https://drive.google.com/drive/folders/1AbCdEfGhIjKlMn...",
        )

        if drive_link:
            fid, fkind = parse_drive_id(drive_link)
            if fid:
                st.caption(f"✅ Detected {fkind} ID: `{fid}`")
            else:
                st.warning("⚠️ Could not detect a valid Drive ID from this link.")

        can_fetch = bool(drive_link.strip()) if drive_link else False

        if st.button("☁️ Fetch & Process from Google Drive", type="primary", disabled=not can_fetch, key="btn_drive"):
            folder_id, kind = parse_drive_id(drive_link)

            if not folder_id:
                st.error("❌ Invalid Google Drive link. Please check and try again.")
            else:
                prog   = st.progress(0)
                status = st.empty()

                raw_files, errors = fetch_resumes_from_drive(
                    folder_id, GOOGLE_DRIVE_API_KEY, status, prog
                )

                prog.empty(); status.empty()

                if errors:
                    for err in errors:
                        st.warning(f"⚠️ {err}")

                if raw_files:
                    all_data: List[Dict] = []
                    failed:   List[str]  = []
                    prog2   = st.progress(0)
                    status2 = st.empty()

                    for i, item in enumerate(raw_files):
                        status2.text(f"Extracting {i + 1}/{len(raw_files)}: {item['filename']}")
                        try:
                            result = process_resume(item['bytes'], item['filename'])
                            if result:
                                all_data.append(result)
                            else:
                                failed.append(item['filename'])
                        except Exception as e:
                            failed.append(f"{item['filename']} ({e})")
                        prog2.progress((i + 1) / len(raw_files))

                    prog2.empty(); status2.empty()

                    if failed:
                        st.warning("⚠️ Skipped: " + ", ".join(failed[:5]))

                    st.session_state.processed_data  = all_data
                    st.session_state.processing_done = True
                    st.success(f"✅ Processed **{len(all_data)}** resumes from Google Drive!")
                    st.rerun()

# ════════════════════════════════════════════════════════════════════
# STEP 2 — SEARCH
# ════════════════════════════════════════════════════════════════════
else:
    data: List[Dict] = st.session_state.processed_data
    st.success(f"✅ **{len(data)}** resumes loaded and ready to search")

    st.divider()
    st.header("② Search & Filter")

    col_input, col_mode = st.columns([4, 1])
    with col_input:
        search_input: str = st.text_input(
            "🔍 Enter search terms — comma-separated",
            value=st.session_state.last_terms,
            placeholder="e.g.,  IT, Pharma, Tata Consultancy, FMCG, Sales, Bangalore, Python, 5 years, MBA",
        )
    with col_mode:
        st.write(""); st.write("")
        match_mode: str = st.radio(
            "Match mode",
            ["Any term (OR)", "All terms (AND)"],
            index=0,
        )

    search_clicked = st.button("🔍 Search Resumes", type="primary", disabled=not search_input.strip())

    if search_clicked and search_input.strip():
        terms     = [t.strip() for t in search_input.split(',') if t.strip()]
        match_all = "AND" in match_mode
        m, nm     = search_resumes(data, terms, match_all)
        st.session_state.matching     = m
        st.session_state.non_matching = nm
        st.session_state.searched     = True
        st.session_state.last_terms   = search_input

    # ── STEP 3 — RESULTS ─────────────────────────────────────────
    if st.session_state.searched:
        matching:     List[Dict] = st.session_state.matching
        non_matching: List[Dict] = st.session_state.non_matching
        all_terms_used = [t.strip() for t in st.session_state.last_terms.split(',') if t.strip()]

        st.divider()
        st.header("③ Results")
        st.caption(
            f"Search terms: **{', '.join(all_terms_used)}**  |  "
            f"Mode: **{match_mode}**  |  "
            f"✅ {len(matching)} matching  ·  ❌ {len(non_matching)} non-matching"
        )

        st.subheader(f"✅ Matching Resumes — {len(matching)}")
        if matching:
            st.dataframe(pd.DataFrame(matching), use_container_width=True, height=min(400, 80 + len(matching) * 38))
            dl_col1, _ = st.columns([1, 3])
            with dl_col1:
                st.download_button(
                    "📥 Download Matching (Excel)",
                    data=to_excel(matching),
                    file_name="matching_resumes.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary",
                    key="dl_matching",
                )
        else:
            st.info("No resumes matched. Try different keywords or switch to **Any (OR)** mode.")

        st.divider()

        st.subheader(f"❌ Non-Matching Resumes — {len(non_matching)}")
        if non_matching:
            with st.expander(f"Show {len(non_matching)} non-matching resumes", expanded=False):
                st.dataframe(pd.DataFrame(non_matching), use_container_width=True, height=min(400, 80 + len(non_matching) * 38))
                st.download_button(
                    "📥 Download Non-Matching (Excel)",
                    data=to_excel(non_matching),
                    file_name="non_matching_resumes.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_non_matching",
                )
        else:
            st.success("🎉 Every resume matched your search terms!")

        st.divider()
        st.download_button(
            "📥 Download All Resumes (Excel)",
            data=to_excel(matching + non_matching),
            file_name="all_resumes.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_all",
        )

    else:
        st.divider()
        st.subheader("📊 All Processed Resumes (search above to filter)")
        display = [{k: v for k, v in d.items() if k != '_full_text'} for d in data]
        st.dataframe(pd.DataFrame(display), use_container_width=True, height=450)
        st.download_button(
            "📥 Download All (Excel)",
            data=to_excel(display),
            file_name="all_resumes.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_all_initial",
        )

# ── Footer ─────────────────────────────────────────────────────────
st.divider()
st.caption("Have a GOOD DAY!!!")